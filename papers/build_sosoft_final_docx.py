#!/usr/bin/env python3
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "sosoft-keymaps-doc.md"
OUTPUT = ROOT / "sosoft-keymaps-final-with-edit-log.docx"
ARTICLE_URL = "https://sosoft.arts.ucla.edu/keymaps-as-social-software/"
RECEIPT_URL = "https://docs.google.com/document/d/1SaHuqqetIFoDhfx3YT3aqoxYSBSh2VVHLhBjZAO3_u4/edit"

CASEY_CHANGES = [
    "Removed the “carry it like a tune” flourish and stated the portability claim directly.",
    "Replaced the ambiguous use of “program” with the Social Software initiative at UCLA Design Media Arts.",
    "Identified Ableton Live, GarageBand, and Logic Pro as audio apps for general readers.",
    "Changed “it eats both hands” to the literal “requires both hands.”",
    "Removed the confusing desire-path comparison and described repetition and learned inertia instead.",
    "Added QWERTY row stagger, handedness, and chord shapes to the AWSED ergonomics critique.",
    "Introduced Vim as the Vim text editor at first mention.",
    "Replaced the unclear “music-app staircase” back-reference with AWSED by name.",
    "Changed “And not only people read…” to “People aren’t the only ones reading…”",
    "Changed “This page” to “This essay.”",
    "Explained how the edition itself acts as social software through a shared printing and assembly agreement.",
    "Added a direct prompt inviting readers to press the letter keys in the interactive keyboard.",
]


def hyperlink(paragraph, text, url):
    rel_id = paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.extend((color, underline))
    run.append(props)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    link.append(run)
    paragraph._p.append(link)


TOKEN = re.compile(r"(\[([^\]]+)\]\((https?://[^)]+)\)|\*\*([^*]+)\*\*|\*([^*]+)\*|`([^`]+)`)")


def inline(paragraph, text):
    at = 0
    for match in TOKEN.finditer(text):
        if match.start() > at:
            paragraph.add_run(text[at:match.start()])
        if match.group(2):
            hyperlink(paragraph, match.group(2), match.group(3))
        elif match.group(4):
            paragraph.add_run(match.group(4)).bold = True
        elif match.group(5):
            paragraph.add_run(match.group(5)).italic = True
        elif match.group(6):
            run = paragraph.add_run(match.group(6))
            run.font.name = "Menlo"
        at = match.end()
    if at < len(text):
        paragraph.add_run(text[at:])


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = section.bottom_margin = Inches(1)
section.left_margin = section.right_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(11)
styles["Title"].font.name = "Arial"
styles["Title"].font.size = Pt(24)
for name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
    styles[name].font.name = "Arial"
    styles[name].font.size = Pt(size)
    styles[name].font.color.rgb = RGBColor(0, 0, 0)

lines = SOURCE.read_text().splitlines()
for line in lines:
    if line.startswith("# "):
        doc.add_heading(line[2:], level=0)
    elif line.startswith("## "):
        doc.add_heading(line[3:], level=1)
    elif not line.strip():
        continue
    else:
        p = doc.add_paragraph()
        inline(p, line)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.1

doc.add_page_break()
doc.add_heading("Casey Reas copy-edit log", level=1)
p = doc.add_paragraph("All twelve anchored comments in the review draft are incorporated in the source above. The detailed repository log records the exact before/after language; this document keeps the review-level summary.")
p.paragraph_format.space_after = Pt(10)
for item in CASEY_CHANGES:
    doc.add_paragraph(item, style="List Number")

doc.add_heading("Publication and review links", level=1)
p = doc.add_paragraph()
hyperlink(p, "Published article — The Keymap Is the Score", ARTICLE_URL)
p = doc.add_paragraph()
hyperlink(p, "Revision 02 video receipt and screenplay", RECEIPT_URL)

doc.save(OUTPUT)
print(OUTPUT)
