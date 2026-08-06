#!/usr/bin/env python3
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
REV = ROOT / "revision-02"
QA = ROOT.parent / "qa"
OUTPUT = REV / "scores-for-social-software-r02-receipt.docx"
REVIEW_VIDEO = "https://drive.google.com/file/d/17jMTkVX7OCekLxKq8X1-4iRLrWzMA0E5/view"
ARTICLE = "https://sosoft.arts.ucla.edu/keymaps-as-social-software/"
FINAL_SOURCE = "https://docs.google.com/document/d/1hNzUm3SmsEBRtM3zWhcQqsYvsoRf4ZioFIQMFndlwXY/edit"

BEATS = [
    ("SSF-00", "Introduction", "00:00–00:18", "SSF-00-introduction.png", "This blue folder just arrived from Social Software at UCLA. Inside is Scores for Social Software, a spring 2026 edition of sixty-four. The copy in my hands is number fifty-one, assembled after ten weeks of making, testing, and performing together."),
    ("SSF-01", "Jeffrey Alan Scudder — Notepat", "00:18–00:31", "SSF-01-notepat.png", "My contribution is Notepat. A folded white user manual opens into a pointed shape, combining an illustrated player, a QR code, instructions, and circular keyboard diagrams."),
    ("SSF-02", "Æther Cavendish — Vigil Score", "00:31–00:38", "SSF-02-vigil-score.png", "Æther Cavendish’s Vigil Score arrives as a matte-black folded packet, closed with a small circular silver seal."),
    ("SSF-03", "Chelly Jin — Software as a Choreography", "00:38–00:48", "SSF-03-software-as-choreography.png", "Chelly Jin’s Software as a Choreography presents a grid of silhouetted hands and arms, pairing each gesture with marks for timing and movement."),
    ("SSF-04", "Jordan Silver — Sonic Architecture", "00:48–01:01", "SSF-04-sonic-architecture.png", "Jordan Silver’s Sonic Architecture is a single sheet of printed columns, large numbers, and typewritten commands, anchored by a looping spiral diagram for inhabiting and measuring space through sound."),
    ("SSF-05", "Em Lugo — Cues for Losing Direction", "01:01–01:11", "SSF-05-cues-for-losing-direction.png", "Em Lugo’s Cues for Losing Direction fits on a small black card, its title set in pale condensed type like a portable instruction carried in the hand."),
    ("SSF-06", "Darlyn Phan — Line Piece 1", "01:11–01:22", "SSF-06-line-piece-1.png", "Darlyn Phan’s Line Piece 1 begins on translucent white paper. Its faint rainbow cast and minimal title let the page behave like a line or veil."),
    ("SSF-07", "Thomas Noya — Biophonía", "01:22–01:34", "SSF-07-biophonia.png", "Thomas Noya’s Biophonía is a branching field of blue, yellow, black, and beige organic forms. In the video, these cell-like clusters drift and reorganize."),
    ("SSF-08", "Banyi Huang — A Cosmographic Score…", "01:34–01:44", "SSF-08-cosmographic-score.png", "Banyi Huang’s A Cosmographic Score for Folding Back into the Kernel centers a luminous circle and a small diagram of connected nodes inside a lavender field."),
    ("SSF-09", "Alexander Espinosa — Music for World Computers", "01:44–01:57", "SSF-09-music-for-world-computers.png", "Alexander Espinosa’s Music for World Computers is a white typographic score: awake, expand, decrease, oxygen, forest, junk, cinnamon, hammer."),
    ("SSF-10", "Mavyn Vu — The Radio Is an Altar: Portal", "01:57–02:07", "SSF-10-radio-altar-portal.png", "Mavyn Vu’s The Radio Is an Altar: Portal combines translucent blue and white score cards, a target-like radio image, small figures, and instructions arranged around their edges."),
    ("SSF-11", "Closing", "02:07–02:21", "SSF-11-closing.png", "Casey Reas facilitated the cycle, with Lauren Lee McCarthy and the Social Software community. Together, the contributions open many paths through a question: if software organizes behavior, what else can we ask it to organize?"),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="C8C8C8", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.append(color)
    props.append(underline)
    run.append(props)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    link.append(run)
    paragraph._p.append(link)


def add_scene(doc, beat, page_break_before=False):
    scene_id, title, timecode, frame, narration = beat
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(5)
    heading.paragraph_format.space_after = Pt(5)
    heading.paragraph_format.keep_with_next = True
    heading.paragraph_format.page_break_before = page_break_before
    r = heading.add_run(f"{scene_id} · {title}")
    r.bold = True
    r.font.size = Pt(14)
    heading.add_run(f"   {timecode}").font.size = Pt(10)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    left, right = table.rows[0].cells
    left.width = Inches(2.05)
    right.width = Inches(4.25)
    for cell in (left, right):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_border(cell)
    left.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    left.paragraphs[0].add_run().add_picture(str(QA / frame), width=Inches(1.72))
    right.paragraphs[0].add_run("Narration").bold = True
    p = right.add_paragraph(narration)
    p.paragraph_format.space_after = Pt(7)
    p = right.add_paragraph()
    p.add_run("Requested changes").bold = True
    request = right.add_paragraph(f"Add a comment here and name {scene_id}.")
    request.paragraph_format.space_after = Pt(3)
    set_cell_shading(right, "FFF8D8")


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(11)
styles["Title"].font.name = "Arial"
styles["Title"].font.size = Pt(24)
for name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
    styles[name].font.name = "Arial"
    styles[name].font.size = Pt(size)
    styles[name].font.color.rgb = RGBColor(0, 0, 0)

doc.add_heading("Scores for Social Software", 0)
p = doc.add_paragraph()
p.add_run("Revision 02 video receipt + timecoded screenplay").bold = True
p = doc.add_paragraph("Review cut · 02:21 · 1080 × 1920 vertical · prepared 27 July 2026")
p.style = styles["Normal"]
p.runs[0].font.color.rgb = RGBColor(90, 90, 90)

link_box = doc.add_table(rows=3, cols=1)
link_box.alignment = WD_TABLE_ALIGNMENT.LEFT
for cell in link_box.column_cells(0):
    set_cell_shading(cell, "FFF8D8")
    set_cell_border(cell, "EAD88A")
add_hyperlink(link_box.cell(0, 0).paragraphs[0], "Watch the revision 02 review video", REVIEW_VIDEO)
add_hyperlink(link_box.cell(1, 0).paragraphs[0], "Open the final article source + Casey edit log", FINAL_SOURCE)
add_hyperlink(link_box.cell(2, 0).paragraphs[0], "Open the published article: The Keymap Is the Score", ARTICLE)

doc.add_heading("How to review", level=1)
doc.add_paragraph("Comment on the words or image you want changed and include the stable scene ID—for example, SSF-04. Scene IDs persist even when later edits move the timestamps.")

doc.add_heading("What changed since revision 01", level=1)
for text in (
    "SSF-02: corrected Æther Cavendish’s Vigil Score to the matte-black folded packet visible in the film.",
    "SSF-04: replaced the provisional Jordan Silver description with the observed single-sheet layout, printed columns, large numbers, typewritten commands, and spiral diagram.",
    "SSF-11: replaced “every contribution asks the same question” with a close that preserves the distinct entry offered by every contribution.",
    "Re-timed all scenes and rebuilt the lower information panel with exact captions, current artist/work labels, and a chapter timeline.",
):
    doc.add_paragraph(text, style="List Bullet")

screenplay_heading = doc.add_heading("Timecoded screenplay", level=1)
screenplay_heading.paragraph_format.page_break_before = True
for index, beat in enumerate(BEATS):
    add_scene(doc, beat, page_break_before=index > 0)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Revision chain: ").bold = True
p.add_run("comment → accepted note → revised scene → new receipt. Accepted comments are copied to accepted-feedback.md under the same scene ID before revision 03 begins.")

for paragraph in doc.paragraphs:
    paragraph.paragraph_format.line_spacing = 1.08

REV.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
